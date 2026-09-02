"""DOCX 解析器"""
from pathlib import Path
from app.parsers.base import BaseParser, ParsedDocument, ParsedSection
from app.utils.logging import log


class DOCXParser(BaseParser):
    """Word .docx 解析 - 使用 python-docx"""

    def parse(self, file_path: str | Path) -> ParsedDocument:
        from docx import Document

        path = Path(file_path)
        log.info(f"开始解析 DOCX: {path.name}")

        doc = Document(str(path))
        sections = []
        full_text_parts = []
        current_section = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            full_text_parts.append(text)

            # 判断标题
            style_name = para.style.name if para.style else ""
            if "Heading" in style_name or "标题" in style_name:
                # 新章节
                if current_section:
                    sections.append(ParsedSection(
                        title=current_section[0],
                        text="\n".join(current_section[1:]),
                        level=1,
                    ))
                current_section = [text]
            else:
                current_section.append(text)

        # 最后一个章节
        if current_section:
            sections.append(ParsedSection(
                title=current_section[0] if current_section[0] != "" else path.stem,
                text="\n".join(current_section[1:] if current_section[0] else current_section),
                level=1,
            ))

        # 如果没有标题结构，把整个文档作为一个 section
        if not sections:
            full_text = "\n".join(full_text_parts)
            sections = [ParsedSection(
                title=path.stem,
                text=full_text,
                level=1,
            )]

        full_text = "\n\n".join(full_text_parts)
        metadata = {
            "file_name": path.name,
            "file_type": "docx",
            "paragraphs_count": len(doc.paragraphs),
        }

        # core properties
        try:
            cp = doc.core_properties
            if cp:
                if cp.title:
                    metadata["title"] = cp.title
                if cp.author:
                    metadata["author"] = cp.author
        except Exception:
            pass

        log.info(f"DOCX 解析完成: {len(sections)} 章节, {len(full_text)} 字符")
        return ParsedDocument(full_text=full_text, sections=sections, metadata=metadata)
